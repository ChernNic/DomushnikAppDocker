import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from graphviz import Digraph
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import sys
import threading

class ProjectAnalyzerApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Project Analyzer")
        self.master.geometry("700x700")

        self.project_path = ""
        self.selected_folders = []
        self.selected_extensions = []
        self.output_path = ""

        self.create_widgets()

    def create_widgets(self):
        # Step 1: Select Project Folder
        self.step1_frame = ttk.LabelFrame(self.master, text="1. Выбор корневой папки проекта")
        self.step1_frame.pack(pady=10, padx=10, fill=tk.X)

        self.select_folder_button = ttk.Button(self.step1_frame, text="Выбрать папку проекта", command=self.select_project_folder)
        self.select_folder_button.pack(pady=5)

        self.project_label = ttk.Label(self.step1_frame, text="Папка проекта не выбрана")
        self.project_label.pack(pady=5)

        # Step 2: Select Subfolders
        self.step2_frame = ttk.LabelFrame(self.master, text="2. Выбор подкаталогов проекта")
        # Initially hidden

        self.subfolders_listbox = tk.Listbox(self.step2_frame, selectmode=tk.MULTIPLE, height=10)
        self.subfolders_listbox.pack(pady=5, padx=5, fill=tk.BOTH, expand=True)

        self.confirm_subfolders_button = ttk.Button(self.step2_frame, text="Подтвердить выбор папок", command=self.confirm_subfolders_selection)
        self.confirm_subfolders_button.pack(pady=5)

        # Step 3: Select File Extensions
        self.step3_frame = ttk.LabelFrame(self.master, text="3. Выбор расширений файлов")
        # Initially hidden

        self.extensions_listbox = tk.Listbox(self.step3_frame, selectmode=tk.MULTIPLE, height=10)
        self.extensions_listbox.pack(pady=5, padx=5, fill=tk.BOTH, expand=True)

        self.confirm_extensions_button = ttk.Button(self.step3_frame, text="Подтвердить выбор расширений", command=self.confirm_extensions_selection)
        self.confirm_extensions_button.pack(pady=5)

        # Step 4: Select Output Path
        self.step4_frame = ttk.LabelFrame(self.master, text="4. Выбор пути для сохранения документа")
        # Initially hidden

        self.select_output_button = ttk.Button(self.step4_frame, text="Выбрать путь для сохранения docx", command=self.select_output_path)
        self.select_output_button.pack(pady=5)

        self.output_label = ttk.Label(self.step4_frame, text="Путь для сохранения не выбран")
        self.output_label.pack(pady=5)

        # Progress Bar
        self.progress_frame = ttk.Frame(self.master)
        self.progress = ttk.Progressbar(self.progress_frame, mode='indeterminate')
        self.progress.pack(fill=tk.X, padx=10, pady=10)

    def select_project_folder(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.project_path = folder_selected
            self.project_label.config(text=f"Выбранная папка: {self.project_path}")
            self.populate_subfolders()
            # Show step 2
            self.step2_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    def populate_subfolders(self):
        self.subfolders_listbox.delete(0, tk.END)
        subfolders = [f.name for f in os.scandir(self.project_path) if f.is_dir()]
        for folder in subfolders:
            self.subfolders_listbox.insert(tk.END, folder)

    def confirm_subfolders_selection(self):
        selected_indices = self.subfolders_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите хотя бы одну папку.")
            return
        self.selected_folders = [self.subfolders_listbox.get(i) for i in selected_indices]
        self.populate_extensions()
        # Show step 3
        self.step3_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)

    def populate_extensions(self):
        self.extensions_listbox.delete(0, tk.END)
        extensions = set()
        for folder in self.selected_folders:
            folder_path = os.path.join(self.project_path, folder)
            for root, dirs, files in os.walk(folder_path):
                for file in files:
                    if file.endswith(".g.dart"):
                        continue  # Игнорируем файлы с расширением .g.dart
                    ext = os.path.splitext(file)[1]
                    if ext:
                        extensions.add(ext)
        for ext in sorted(extensions):
            self.extensions_listbox.insert(tk.END, ext)

    def confirm_extensions_selection(self):
        selected_indices = self.extensions_listbox.curselection()
        if not selected_indices:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите хотя бы одно расширение.")
            return
        self.selected_extensions = [self.extensions_listbox.get(i) for i in selected_indices]
        # Show step 4
        self.step4_frame.pack(pady=10, padx=10, fill=tk.BOTH, expand=True)
        # Automatically invoke the output path selection
        self.select_output_path()

    def select_output_path(self):
        file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                 filetypes=[("Word Document", "*.docx")])
        if file_path:
            self.output_path = file_path
            self.output_label.config(text=f"Путь сохранения: {self.output_path}")
            # Start generation
            self.start_generation()

    def start_generation(self):
        # Start generation in a separate thread to keep the GUI responsive
        threading.Thread(target=self.generate_docx, daemon=True).start()

    def generate_docx(self):
        try:
            # Show progress bar
            self.progress_frame.pack(pady=10, padx=10, fill=tk.X)
            self.progress.start(10)

            # Collect project structure
            project_structure = {}
            for folder in self.selected_folders:
                folder_path = os.path.join(self.project_path, folder)
                self.add_to_structure(project_structure, folder, folder_path)

            # Collect file information
            modules = []
            module_details = []
            code_contents = {}
            for idx, (module_path, full_path) in enumerate(self.get_files(project_structure), 1):
                rel_path = os.path.relpath(full_path, self.project_path)
                modules.append((idx, f"/{rel_path.replace(os.sep, '/')}"))
                # Get number of lines and size
                try:
                    with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                        lines = f.readlines()
                        num_lines = len(lines)
                        code = ''.join(lines)
                except Exception as e:
                    lines = []
                    num_lines = 0
                    code = f"Не удалось прочитать файл: {e}"
                size_kb = round(os.path.getsize(full_path) / 1024, 2)
                module_details.append((f"/{rel_path.replace(os.sep, '/')}", "", num_lines, size_kb))
                code_contents[f"/{rel_path.replace(os.sep, '/')}"] = code

            # Generate graph
            graph_image_path = os.path.join(os.path.dirname(sys.argv[0]), "project_structure")
            self.create_graph(project_structure, graph_image_path)

            # Create docx
            document = Document()

            # Set default styles
            style = document.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14)
            style.paragraph_format.left_indent = Cm(1.25)
            style.paragraph_format.line_spacing = 1.5
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            # Add graph image
            document.add_picture(graph_image_path + ".png", width=Pt(500))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add first table
            document.add_paragraph()  # Add a blank paragraph
            table1_title = document.add_paragraph("Структурная схема проекта")
            table1_title.style.font.size = Pt(12)
            table1_title.style.font.name = 'Times New Roman'
            table1_title.paragraph_format.left_indent = Cm(0)
            table1_title.paragraph_format.line_spacing = 1
            table1_title.paragraph_format.space_before = Pt(0)
            table1_title.paragraph_format.space_after = Pt(0)
            table1_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            table1 = document.add_table(rows=1, cols=3)
            table1.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table1.rows[0].cells
            hdr_cells[0].text = 'Номер'
            hdr_cells[1].text = 'Название модуля'
            hdr_cells[2].text = 'Описание'

            for num, mod in modules:
                row_cells = table1.add_row().cells
                row_cells[0].text = str(num)
                row_cells[1].text = mod
                row_cells[2].text = ""

            # Add second table
            document.add_paragraph()  # Add a blank paragraph
            table2_title = document.add_paragraph("Модульная информация")
            table2_title.style.font.size = Pt(12)
            table2_title.style.font.name = 'Times New Roman'
            table2_title.paragraph_format.left_indent = Cm(0)
            table2_title.paragraph_format.line_spacing = 1
            table2_title.paragraph_format.space_before = Pt(0)
            table2_title.paragraph_format.space_after = Pt(0)
            table2_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            table2 = document.add_table(rows=1, cols=4)
            table2.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = table2.rows[0].cells
            hdr_cells[0].text = 'Модуль'
            hdr_cells[1].text = 'Описание'
            hdr_cells[2].text = 'Количество строк'
            hdr_cells[3].text = 'Размер (Кб)'

            for detail in module_details:
                row_cells = table2.add_row().cells
                row_cells[0].text = detail[0]
                row_cells[1].text = detail[1]
                row_cells[2].text = str(detail[2])
                row_cells[3].text = str(detail[3])

            # Add code sections
            for mod_path, code in code_contents.items():
                document.add_paragraph()  # Blank paragraph
                code_title = document.add_paragraph(mod_path)
                code_title.style.font.size = Pt(12)
                code_title.style.font.name = 'Times New Roman'
                code_title.paragraph_format.left_indent = Cm(0)
                code_title.paragraph_format.line_spacing = 1
                code_title.paragraph_format.space_before = Pt(0)
                code_title.paragraph_format.space_after = Pt(0)
                code_title.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                code_paragraph = document.add_paragraph(code)
                code_paragraph.style.font.size = Pt(9.5)
                code_paragraph.style.font.name = 'Times New Roman'
                code_paragraph.paragraph_format.left_indent = Cm(0)
                code_paragraph.paragraph_format.line_spacing = 1
                code_paragraph.paragraph_format.space_before = Pt(0)
                code_paragraph.paragraph_format.space_after = Pt(0)
                code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Save document
            document.save(self.output_path)

            # Stop progress bar
            self.progress.stop()
            self.progress_frame.pack_forget()

            # Notification
            messagebox.showinfo("Успех", f"Документ успешно сгенерирован по пути:\n{self.output_path}")

        except Exception as e:
            # Stop progress bar in case of error
            self.progress.stop()
            self.progress_frame.pack_forget()
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")

    def add_to_structure(self, structure, folder_name, folder_path):
        structure[folder_name] = []
        for item in os.listdir(folder_path):
            item_path = os.path.join(folder_path, item)
            if os.path.isdir(item_path):
                structure[folder_name].append({item: []})
                self.add_to_structure(structure[folder_name][-1], item, item_path)
            else:
                if item.endswith(".g.dart"):
                    continue  # Игнорируем файлы с расширением .g.dart
                _, ext = os.path.splitext(item)
                if ext in self.selected_extensions:
                    structure[folder_name].append(item)

    def get_files(self, structure, parent_path=""):
        for key, value in structure.items():
            current_path = os.path.join(parent_path, key)
            for item in value:
                if isinstance(item, dict):
                    yield from self.get_files(item, current_path)
                else:
                    yield (os.path.join(current_path, item), os.path.join(self.project_path, current_path, item))

    def create_graph(self, structure, output_path, parent=None, graph=None):
        if graph is None:
            graph = Digraph(comment='Project Structure', format='png')

            # Настройки графа для горизонтального расположения
            graph.attr(rankdir='LR')  # Left to Right
            graph.attr(splines='ortho')  # Orthogonal lines

            # Установка качества
            graph.attr(dpi='300')  # High resolution

        for key, value in structure.items():
            current_node = os.path.join(parent, key) if parent else key
            # Удаление обратных слэшей для корректного отображения
            current_node_clean = current_node.replace('\\', '/')
            # Добавление папки
            graph.node(current_node_clean, key, shape='box', style='filled', fillcolor='lightblue')
            if parent:
                parent_clean = parent.replace('\\', '/')
                graph.edge(parent_clean, current_node_clean)
            for item in value:
                if isinstance(item, dict):
                    child_key = list(item.keys())[0]
                    child_node = os.path.join(current_node, child_key)
                    child_node_clean = child_node.replace('\\', '/')
                    graph.node(child_node_clean, child_key, shape='box', style='filled', fillcolor='lightblue')
                    graph.edge(current_node_clean, child_node_clean)
                    self.create_graph(item, output_path, current_node, graph)
                else:
                    file_node = os.path.join(current_node, item)
                    file_node_clean = file_node.replace('\\', '/')
                    graph.node(file_node_clean, item, shape='note', style='filled', fillcolor='lightyellow')
                    graph.edge(current_node_clean, file_node_clean)
        graph.render(output_path, cleanup=True)

def main():
    root = tk.Tk()
    app = ProjectAnalyzerApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
